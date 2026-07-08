"""Word .docx text extractor (python-docx). Deterministic for a fixed lib version.

Registered by the registry only when `python-docx` imports, so a box without it
captures .docx bytes-only (ADR-0010 rule 5) rather than crashing. The old binary
`.doc` format is a different beast and is intentionally not handled here.
"""
from __future__ import annotations

import io

import docx as _docx

from .base import Extractor

try:
    from importlib.metadata import version as _pkg_version
    _VERSION = _pkg_version("python-docx")
except Exception:  # pragma: no cover
    _VERSION = "?"


class DocxExtractor(Extractor):
    name = f"python-docx@{_VERSION}"
    extensions = frozenset({".docx"})

    def extract(self, raw: bytes) -> str:
        doc = _docx.Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        # Flatten table cells too (tab-separated rows) — lossy, but the canonical
        # .docx bytes retain the real structure (ADR-0010).
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
